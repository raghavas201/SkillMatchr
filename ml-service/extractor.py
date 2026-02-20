"""
Text extraction from PDF and DOCX files.
"""

import io
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.layout import LAParams
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF byte buffer."""
    try:
        laparams = LAParams(
            line_margin=0.5,
            word_margin=0.1,
            char_margin=2.0,
            all_texts=True,
        )
        text = pdf_extract_text(
            io.BytesIO(file_bytes),
            laparams=laparams,
        )
        return (text or "").strip()
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {e}") from e


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX byte buffer."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {e}") from e


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Route to the correct extractor based on file type."""
    if file_type.lower() == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif file_type.lower() == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
